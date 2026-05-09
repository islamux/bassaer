import os
import re
from docx import Document

DOCX_PATH = 'content/ar-basaar.docx'
CHAPTERS_DIR = 'content/chapters'

MAPPINGS = {
    'intro.md': (37, 81),
    'chapter-1.md': (82, 1318),
    'chapter-2.md': (1319, 1704),
    'chapter-3.md': (1705, 2857),
    'chapter-4.md': (2858, 3188),
    'chapter-5.md': (3189, 3457),
    'chapter-6.md': (3458, 3685),
    'chapter-7.md': (3686, 4238),
    'chapter-8.md': (4239, 4566),
    'chapter-9.md': (4567, 7371),
    'chapter-10.md': (7372, 9855),
    'chapter-11.md': (9856, 10099),
    'chapter-12.md': (10100, 10500)
}

def clean_asterisks(text):
    text = re.sub(r'\*\*\s+', ' **', text)
    text = re.sub(r'\s+\*\*', '** ', text)
    text = text.replace('****', '')
    text = re.sub(r'\*\*\((\d*)\)\*\*', r'(\1)', text)
    return text.strip()

def run_to_md(run):
    text = run.text
    if not text.strip():
        return text
    if run.bold:
        return f"**{text}**"
    return text

def convert_paragraph(p):
    if not p.text.strip():
        return ""
    
    style = p.style.name
    para_text = "".join(run_to_md(r) for r in p.runs)
    para_text = clean_asterisks(para_text)
    
    if 'نمط7' in style:
        return ""
    elif 'نمط12' in style:
        match = re.match(r'^(\*\*)?\s*\d+\s*-\s*(.*)', para_text)
        if match:
            clean_title = match.group(2)
            has_bold = match.group(1)
            if has_bold and not clean_title.startswith('**'):
                clean_title = f"**{clean_title}"
        else:
            clean_title = para_text
        return f"## {clean_title}"
    else:
        return para_text

def get_frontmatter_and_h1(filepath):
    if not os.path.exists(filepath):
        return "", ""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    frontmatter, h1 = "", ""
    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            frontmatter = f"---{parts[1]}---\n"
            content = parts[2]
    for line in content.split('\n'):
        if line.strip().startswith('# '):
            h1 = line.strip() + '\n\n'
            break
    return frontmatter, h1

def process_file(filename, start_idx, end_idx, doc):
    filepath = os.path.join(CHAPTERS_DIR, filename)
    frontmatter, h1 = get_frontmatter_and_h1(filepath)
    if not h1 and filename == 'intro.md':
        h1 = "# مقدمة الكتاب\n\n"
        
    lines = []
    h2_counter = 1
    
    for i in range(start_idx, end_idx + 1):
        if i >= len(doc.paragraphs):
            break
        
        md_text = convert_paragraph(doc.paragraphs[i])
        if md_text:
            if md_text.startswith('## '):
                title = md_text[3:].strip()
                if filename != 'intro.md':
                    if title.startswith('**'):
                        md_text = f"## {h2_counter} - {title}"
                    else:
                        md_text = f"## {h2_counter} - {title}"
                    h2_counter += 1
                else:
                    md_text = f"## {title}"
            lines.append(md_text)
            
    final_content = frontmatter + h1 + "\n\n".join(lines) + "\n"
    
    # Backup existing
    if os.path.exists(filepath):
        os.rename(filepath, filepath + '.bak')
        
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(final_content)
        
    print(f"Generated {filepath} ({len(lines)} blocks)")

if __name__ == '__main__':
    import sys
    chapters = sys.argv[1:] if len(sys.argv) > 1 else list(MAPPINGS.keys())
    for filename in chapters:
        if filename not in MAPPINGS:
            print(f"Unknown chapter: {filename}. Valid: {list(MAPPINGS.keys())}")
            sys.exit(1)
        start, end = MAPPINGS[filename]
        print(f"[{filename}] Loading DOCX...", flush=True)
        doc = Document(DOCX_PATH)
        process_file(filename, start, end, doc)
        print(f"[{filename}] Done.", flush=True)

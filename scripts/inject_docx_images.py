import os
import re
import sys
import zipfile
from lxml import etree
from docx import Document

DOCX_PATH = 'content/ar-basaar.docx'
CHAPTERS_DIR = 'content/chapters'
IMAGES_OUT_DIR = 'web/public/images'

MAPPINGS = {
    'intro.md': (37, 81),
    'chapter-1.md': (82, 1318),
    'chapter-2.md': (1319, 1704),
    'chapter-3.md': (1705, 2857),
    'chapter-4.md': (2858, 3188),
    'chapter-5.md': (3189, 3457),
    'chapter-6.md': (3458, 3685),
    'chapter-7.md': (3686, 4238),
    'chapter-8.md': (4239, 4566),
    'chapter-9.md': (4567, 7371),
    'chapter-10.md': (7372, 9855),
    'chapter-11.md': (9856, 10099),
    'chapter-12.md': (10100, 10500),
}

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

def build_rid_to_filename():
    with zipfile.ZipFile(DOCX_PATH) as z:
        rels_xml = z.read('word/_rels/document.xml.rels')
        root = etree.fromstring(rels_xml)
        rel_ns = '{http://schemas.openxmlformats.org/package/2006/relationships}'
        rid_map = {}
        for rel in root:
            rid = rel.get('Id')
            target = rel.get('Target')
            if target and target.startswith('media/'):
                rid_map[rid] = target
        return rid_map

def get_para_images(para):
    drawings = para._element.findall('.//w:drawing', NAMESPACES)
    if not drawings:
        return []
    results = []
    for drawing in drawings:
        blip = drawing.find('.//a:blip', NAMESPACES)
        if blip is not None:
            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed and embed in rid_map:
                results.append(rid_map[embed])
    return results

def clean_asterisks(text):
    text = re.sub(r'\*\*\s+', ' **', text)
    text = re.sub(r'\s+\*\*', '** ', text)
    text = text.replace('****', '')
    text = re.sub(r'\*\*\((\d*)\)\*\*', r'(\1)', text)
    return text.strip()

def run_to_md(run):
    text = run.text
    if not text.strip():
        return text
    if run.bold:
        return f'**{text}**'
    return text

def convert_paragraph(p):
    if not p.text.strip():
        return ''
    style = p.style.name
    para_text = ''.join(run_to_md(r) for r in p.runs)
    para_text = clean_asterisks(para_text)
    if 'نمط7' in style:
        return ''
    elif 'نمط12' in style:
        match = re.match(r'^(\*\*)?\s*\d+\s*-\s*(.*)', para_text)
        if match:
            clean_title = match.group(2)
            has_bold = match.group(1)
            if has_bold and not clean_title.startswith('**'):
                clean_title = f'**{clean_title}'
        else:
            clean_title = para_text
        return f'## {clean_title}'
    else:
        return para_text

def get_frontmatter_and_h1(filepath):
    if not os.path.exists(filepath):
        return '', ''
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    frontmatter, h1 = '', ''
    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            frontmatter = f'---{parts[1]}---\n'
            content = parts[2]
    for line in content.split('\n'):
        if line.strip().startswith('# '):
            h1 = line.strip() + '\n\n'
            break
    return frontmatter, h1

def chapter_id_from_filename(filename):
    return filename.replace('.md', '')

def image_output_dir(chapter_id):
    return os.path.join(IMAGES_OUT_DIR, chapter_id)

def extract_and_inject(filename, start_idx, end_idx, doc, rid_map):
    filepath = os.path.join(CHAPTERS_DIR, filename)
    frontmatter, h1 = get_frontmatter_and_h1(filepath)
    if not h1 and filename == 'intro.md':
        h1 = '# مقدمة الكتاب\n\n'

    chapter_id = chapter_id_from_filename(filename)
    out_dir = image_output_dir(chapter_id)
    os.makedirs(out_dir, exist_ok=True)

    lines = []
    h2_counter = 1
    img_counter = 1
    seen_images = set()

    for i in range(start_idx, end_idx + 1):
        if i >= len(doc.paragraphs):
            break

        p = doc.paragraphs[i]
        para_imgs = get_para_images(p)
        md_text = convert_paragraph(p)

        if md_text:
            if md_text.startswith('## '):
                title = md_text[3:].strip()
                if filename != 'intro.md':
                    md_text = f'## {h2_counter} - {title}'
                    h2_counter += 1
                else:
                    md_text = f'## {title}'
            lines.append(md_text)

        # Extract and inject images (even for empty paragraphs with images)
        for img_path in para_imgs:
            if img_path in seen_images:
                continue
            seen_images.add(img_path)

            ext = os.path.splitext(img_path)[1]
            new_name = f'img-{img_counter:03d}{ext}'
            img_counter += 1

            zip_img_path = f'word/{img_path}'
            with zipfile.ZipFile(DOCX_PATH) as z:
                if zip_img_path in z.namelist():
                    data = z.read(zip_img_path)
                    out_path = os.path.join(out_dir, new_name)
                    with open(out_path, 'wb') as f:
                        f.write(data)

            lines.append(f'![](/images/{chapter_id}/{new_name})')

    # Write the markdown with images
    final_content = frontmatter + h1 + '\n\n'.join(lines) + '\n'

    if os.path.exists(filepath):
        os.rename(filepath, filepath + '.bak')

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(final_content)

    print(f'[{filename}] Extracted {img_counter - 1} images, injected into {len(lines)} blocks')

if __name__ == '__main__':
    chapters = sys.argv[1:] if len(sys.argv) > 1 else list(MAPPINGS.keys())
    global rid_map
    rid_map = build_rid_to_filename()

    for filename in chapters:
        if filename not in MAPPINGS:
            print(f'Unknown chapter: {filename}. Valid: {list(MAPPINGS.keys())}')
            sys.exit(1)
        start, end = MAPPINGS[filename]
        print(f'[{filename}] Loading DOCX...', flush=True)
        doc = Document(DOCX_PATH)
        extract_and_inject(filename, start, end, doc, rid_map)
        print(f'[{filename}] Done.', flush=True)
